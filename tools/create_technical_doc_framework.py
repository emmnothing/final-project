from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("docs/technical_documentation_framework.docx")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        add_para(doc, item, style=style)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9.5)
    for name, size in [("Title", 18), ("Heading 1", 13), ("Heading 2", 11)]:
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Technical Documentation Framework")
    run.bold = True
    run.font.size = Pt(18)
    add_para(doc, "Autonomous Mapping and Navigation Robot Project", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Prepared for CDIO Technical Documentation submission | Draft framework before final route freeze", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "Document Control", style="Heading 1")
    add_table(
        doc,
        ["Item", "Draft content"],
        [
            ["Project", "STM32F446-based mobile robot for sensing, mapping, control, and navigation"],
            ["Submission format", "PDF, maximum 10 pages. This DOCX is an editable source draft."],
            ["Current status", "Framework complete; final technical route, final measured data, and final screenshots to be inserted after program freeze."],
            ["Recommended page budget", "Title/control: 0.5 page; CDIO sections: 7 pages; evidence and GenAI: 1.5 pages; conclusion: 0.5 page."],
        ],
        widths=[3.5, 12.5],
    )

    add_para(doc, "1. Executive Summary", style="Heading 1")
    add_para(
        doc,
        "This project develops an embedded mobile robot that senses its environment, estimates its pose, builds a grid-based representation of the surrounding space, and supports controlled driving, autonomous mapping, and navigation experiments. The system is implemented on an STM32F446 microcontroller using FreeRTOS tasks and STM32 HAL peripherals. The technical documentation is structured around the CDIO lifecycle so that each phase shows what was decided, why it was decided, what evidence supports the decision, and how the phase informed the next stage.",
    )
    add_para(
        doc,
        "[TBD after final route] Add one short sentence summarising the final chosen mapping/navigation route and the final operating result.",
        italic=True,
        color=(180, 80, 20),
    )

    add_para(doc, "2. System Context and Stable Architecture", style="Heading 1")
    add_para(
        doc,
        "The stable architecture can already be documented because it is independent of the final tuning route. The robot follows a Sense -> Think -> Act pipeline: sensors and serial inputs produce raw observations; firmware tasks parse, filter, and fuse these observations into map, pose, and command states; motor PWM outputs and serial feedback close the operating loop.",
    )
    add_table(
        doc,
        ["Layer", "Current responsibility", "Evidence to cite"],
        [
            ["Sense", "RPLIDAR scan bytes through USART1 DMA; encoder counts through TIM2/TIM4; MPU6500 through I2C; ADC speed input; Bluetooth commands through USART6.", "Core/Src/lidar_pipeline.c, Core/Src/mpu6500.c, Core/Src/motor_control.c, Core/Src/bluetooth_control.c"],
            ["Think", "FreeRTOS task scheduling, command dispatch, odometry update, grid map update, scan buffering, navigation state, and diagnostic reporting.", "Core/Src/main.c, Core/Src/mapping_grid.c"],
            ["Act", "Differential motor PWM on TIM3, brake/stop modes, forward/turn commands, Bluetooth debug output, OLED pages, and host-side viewers.", "Core/Src/motor_control.c, test.py, test_win_qt.py, test_mac.py"],
        ],
        widths=[2.5, 9.0, 5.0],
    )

    add_para(doc, "3. CDIO Phase 1: Conceive", style="Heading 1")
    add_para(doc, "Purpose", style="Heading 2")
    add_para(
        doc,
        "The Conceive phase translated the open-ended robot challenge into measurable engineering needs: reliable sensing, controlled movement, safe stop behaviour, map feedback, repeatable debugging, and a route for later autonomous operation. The main constraint was that embedded performance and hardware uncertainty had to be managed alongside evolving software functions.",
    )
    add_para(doc, "Evidence already available", style="Heading 2")
    add_bullets(
        doc,
        [
            "Problem interpretation: a small robot must perceive the environment, move under software control, and provide evidence of mapping/navigation progress.",
            "Initial constraints: limited MCU resources, real-time serial data from LIDAR, motor imbalance, sensor noise, Bluetooth reliability, and the need for visible debugging data.",
            "Knowledge gaps identified: LIDAR packet parsing, grid representation, odometry drift, PWM calibration, and validation without relying only on visual observation.",
        ],
    )
    add_para(doc, "Decision justification", style="Heading 2")
    add_para(
        doc,
        "The team treated observability as an early requirement instead of a late testing add-on. This is why the current project contains Bluetooth text diagnostics, map row streaming, odometry debug lines, and host-side visualisers. These choices reduce uncertainty before the final route is frozen.",
    )
    add_para(doc, "Lifecycle link", style="Heading 2")
    add_para(
        doc,
        "This phase informed Design by defining the system as modular pipelines rather than one monolithic control loop: sensing, command handling, mapping, motor actuation, UI, and host-side diagnostics each needed clear boundaries.",
    )

    add_para(doc, "4. CDIO Phase 2: Design", style="Heading 1")
    add_para(doc, "Architecture decisions", style="Heading 2")
    add_table(
        doc,
        ["Design decision", "Reasoning", "Current evidence"],
        [
            ["FreeRTOS task separation", "Separates time-sensitive sensing, UI refresh, and command handling so LIDAR parsing and control feedback do not block each other.", "Task entry points and handles in Core/Src/main.c"],
            ["Dedicated LIDAR pipeline module", "USART1 DMA data is high-volume and benefits from parser/service queues and explicit drop counters.", "Core/Src/lidar_pipeline.c"],
            ["Dedicated Bluetooth command module", "Keeps command parsing, ACK/status output, and USART6 ISR recovery separate from robot logic.", "Core/Src/bluetooth_control.c/h"],
            ["Grid map abstraction", "Allows map insertion/scoring/statistics to be tested and reasoned about separately from motion control.", "Core/Src/mapping_grid.c/h"],
            ["Motor control abstraction", "Encapsulates PWM, encoder state, stop/brake behaviour, and differential drive commands.", "Core/Src/motor_control.c/h"],
        ],
        widths=[4.0, 8.0, 4.0],
    )
    add_para(doc, "Design comparison and route placeholder", style="Heading 2")
    add_para(
        doc,
        "The final report should compare the chosen route with at least one rejected alternative. Keep this concise: one paragraph plus a table is enough. Do not overfill the 10-page limit.",
    )
    add_table(
        doc,
        ["Option", "Advantages", "Limitations", "Final decision"],
        [
            ["[TBD route option A]", "[TBD]", "[TBD]", "[TBD]"],
            ["[TBD route option B]", "[TBD]", "[TBD]", "[TBD]"],
        ],
        widths=[3.5, 4.5, 4.5, 3.5],
    )
    add_para(doc, "Lifecycle link", style="Heading 2")
    add_para(
        doc,
        "The Design phase informed Implementation by defining module boundaries and testable interfaces before adding tuning logic. This helped keep later changes local: for example, command parsing can evolve without rewriting the map grid, and motor tuning can evolve without changing the LIDAR parser.",
    )

    add_para(doc, "5. CDIO Phase 3: Implement", style="Heading 1")
    add_para(doc, "Implementation structure", style="Heading 2")
    add_table(
        doc,
        ["Firmware area", "Implementation status to describe", "Suggested excerpt"],
        [
            ["LIDAR parser", "DMA half/full callbacks queue fixed-size blocks; parser decodes scan nodes and publishes LidarPoint_t data.", "Show LidarPipeline_TakePoint or callback routing."],
            ["Bluetooth control", "Line-based parser turns text/numeric commands into BluetoothCommand_t values and sends ACK/status text.", "Show BluetoothControl_CommandName or parser table."],
            ["Mapping grid", "Occupancy scores represent unknown/free/occupied cells; polar points are projected from pose into the grid.", "Show MappingGrid_InsertPolarPointAtPose or scoring function."],
            ["Motor control", "PWM output and encoder state are encapsulated, with stop/brake and differential drive functions.", "Show MotorControl_SetDifferential or MotorControl_Stop."],
            ["Main application", "Application coordinates mapping state, auto-mapping/navigation state, OLED pages, and debug transmissions.", "Show TestApp_HandleBluetoothCommands and final route function after freeze."],
        ],
        widths=[3.2, 8.5, 4.3],
    )
    add_para(doc, "Code evidence to insert", style="Heading 2")
    add_bullets(
        doc,
        [
            "One short code excerpt for the Sense path: USART1 DMA -> LIDAR point queue.",
            "One short code excerpt for the Think path: map update, scan matching, or final navigation decision logic.",
            "One short code excerpt for the Act path: differential motor command or safe stop path.",
            "One screenshot of the host viewer or serial diagnostic stream showing map/pose/navigation data.",
        ],
    )
    add_para(
        doc,
        "[TBD after final route] Replace the Think-path placeholder with the final selected algorithm/control logic and explain why it is the final version.",
        italic=True,
        color=(180, 80, 20),
    )
    add_para(doc, "Lifecycle link", style="Heading 2")
    add_para(
        doc,
        "Implementation informed Operation by exposing measurable outputs: drop counters, inserted/rejected point counts, pose values, navigation state, obstacle flags, and motor commands. These outputs become the basis for validation instead of relying only on subjective observation.",
    )

    add_para(doc, "6. CDIO Phase 4: Operate", style="Heading 1")
    add_para(doc, "Validation strategy", style="Heading 2")
    add_table(
        doc,
        ["Question", "Measurement / evidence", "Pass criterion"],
        [
            ["Can the robot receive commands reliably?", "Bluetooth ACKs, rx/tx counters, command response logs.", "[TBD final threshold]"],
            ["Can sensing run continuously?", "LIDAR valid/invalid node counts, DMA drop counts, scan start counts.", "[TBD final threshold]"],
            ["Can the robot move controllably?", "Encoder deltas, straight-line drift observation, stop/brake response.", "[TBD final threshold]"],
            ["Can the map update from real scans?", "Map rows, occupied/free cell counts, inserted/rejected point counts.", "[TBD final threshold]"],
            ["Can final navigation/auto-mapping meet the task objective?", "Final route test video/screenshot, timing, success/failure count.", "[TBD after final route]"],
        ],
        widths=[5.0, 7.0, 4.0],
    )
    add_para(doc, "Iteration evidence", style="Heading 2")
    add_para(
        doc,
        "Use this section to show that the system improved through testing. Include two or three compact before/after examples rather than a long diary.",
    )
    add_table(
        doc,
        ["Issue observed", "Evidence", "Refinement", "Result"],
        [
            ["Robot movement too fast during autonomous behaviour", "Observed speed did not match safe mapping operation.", "Autonomous PWM limits separated from manual ADC speed control.", "[TBD measured result]"],
            ["Map/pose reliability needed stronger evidence", "Serial output alone was difficult to inspect.", "Host-side map/pose viewer added for live validation.", "[TBD screenshot]"],
            ["Final route issue", "[TBD]", "[TBD]", "[TBD]"],
        ],
        widths=[4.2, 4.0, 4.5, 3.3],
    )
    add_para(doc, "Lifecycle link", style="Heading 2")
    add_para(
        doc,
        "Operation closes the CDIO loop by feeding measured weaknesses back into design and implementation. The final report should explicitly state which result led to the last refinement and which limitations remain.",
    )

    add_para(doc, "7. GenAI Use and Verification", style="Heading 1")
    add_table(
        doc,
        ["Required example", "Draft content to adapt", "Verification action"],
        [
            ["Useful GenAI output", "GenAI helped convert the marking brief into a CDIO evidence checklist, reminding the team to document lifecycle links, decision justification, and validation evidence rather than only listing final features.", "Checked against the teacher's requirement screenshots and kept as a structure, not as technical proof."],
            ["Unhelpful or misleading GenAI output", "A generated draft tended to sound as if final performance data and final algorithm choices were already confirmed.", "Corrected by marking final route and performance values as TBD until verified by code freeze and measured tests."],
        ],
        widths=[4.0, 8.5, 3.5],
    )
    add_para(
        doc,
        "Responsible use statement: GenAI was used for structuring, summarising options, and checking whether the documentation addresses the marking criteria. Engineering claims, code excerpts, and performance results must be verified against the repository and physical tests before submission.",
    )

    add_para(doc, "8. Evidence Checklist Before PDF Export", style="Heading 1")
    add_table(
        doc,
        ["Evidence item", "Owner", "Status"],
        [
            ["Final system architecture figure or concise table", "[Name]", "Draft table included"],
            ["Final technical route comparison and decision", "[Name]", "TBD"],
            ["Code excerpt: sensing", "[Name]", "TBD"],
            ["Code excerpt: mapping/navigation/control", "[Name]", "TBD"],
            ["Code excerpt: actuation/safety stop", "[Name]", "TBD"],
            ["Serial log or viewer screenshot", "[Name]", "TBD"],
            ["Measured performance table", "[Name]", "TBD"],
            ["GenAI useful and misleading examples", "[Name]", "Draft included"],
            ["Final PDF page count <= 10", "[Name]", "Check after export"],
        ],
        widths=[8.5, 3.0, 4.0],
    )

    add_para(doc, "9. Conclusion", style="Heading 1")
    add_para(
        doc,
        "This documentation framework shows the engineering process rather than only the finished robot. The report will be completed by inserting the final route decision, final code excerpts, and measured operating evidence. Once those final items are added, the document should demonstrate traceability from problem interpretation, through design and implementation, to tested operation.",
    )

    add_para(doc, "Appendix: Source Files Referenced", style="Heading 1")
    add_para(
        doc,
        "Keep this appendix short or remove it if the final PDF exceeds 10 pages. Suggested references: Core/Src/main.c, Core/Src/lidar_pipeline.c, Core/Src/bluetooth_control.c, Core/Src/mapping_grid.c, Core/Src/motor_control.c, Core/Inc/*.h, test.py, test_win_qt.py, test_mac.py.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    make_doc()
    print(OUT)
